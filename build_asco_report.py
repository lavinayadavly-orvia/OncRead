from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/lavinayadav/Documents/ASCO/ASCO_2026_Advances_Impact_Assessment.docx"

NAVY = "123B66"
BLUE = "1676B8"
TEAL = "008C8C"
LIGHT_BLUE = "DDF2FA"
PALE_BLUE = "EDF8FC"
LIGHT_GRAY = "F2F5F8"
MID_GRAY = "596579"
DARK = "1D2A38"
GREEN = "17805C"
GOLD = "B66A00"
RED = "B42318"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run(run, size=11.2, bold=False, color=DARK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=11.2, bold=False, color=DARK, italic=False,
             before=0, after=7, line=1.12, align=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0, color=DARK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run(r, size=11, color=color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 17.5, BLUE, 18, 9
    elif level == 2:
        size, color, before, after = 14, TEAL, 14, 7
    else:
        size, color, before, after = 12.2, NAVY, 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=size, bold=True, color=color)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.35)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=11.4, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    set_run(r2, size=10.9)
    add_para(doc, "", after=2)


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths is None:
        widths = [6.35 / len(headers)] * len(headers)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=9.3, bold=True, color=WHITE)
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i], top=90, start=100, bottom=90, end=100)
            if idx % 2 == 1:
                set_cell_shading(cells[i], LIGHT_BLUE)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.03
            if i > 0 and len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=DARK)
    add_para(doc, "", after=2)
    return table


def add_profile(doc, number, title, status, current_care, design, findings,
                comparative, toxicity, cost, impact, limits, source_nums,
                manufacturer=None, india_access=None):
    add_heading(doc, f"{number}. {title}", 2)
    add_para(doc, f"Evidence status: {status}", size=10.4, bold=True, color=GREEN,
             after=6, keep=True)
    fields = [
        ("Developer / manufacturer", manufacturer),
        ("Current treatment benchmark", current_care),
        ("Study design and population", design),
        ("Measured findings", findings),
        ("Comparison with current care", comparative),
        ("Safety and treatment burden", toxicity),
        ("Cost and access implications", cost),
        ("India access, price and launch outlook", india_access),
        ("Expected impact on cancer care", impact),
        ("Limitations and adoption threshold", limits),
    ]
    for label, text in fields:
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.keep_together = True
        r1 = p.add_run(f"{label}: ")
        set_run(r1, size=10.9, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run(r2, size=10.9)
    add_para(doc, f"Sources: {source_nums}", size=9.3, color=MID_GRAY, after=8)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11.2)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.12

for style_name in ("List Bullet", "List Number"):
    st = styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("ASCO 2026 | Oncology Advances and Impact Assessment")
set_run(hr, size=9, bold=True, color=TEAL)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Evidence review based on ASCO/JCO email alerts received May 20-June 8, 2026")
set_run(fr, size=8.5, color=MID_GRAY)

# Cover
add_para(doc, "EVIDENCE AND TECHNOLOGY REPORT", size=11, bold=True, color=GOLD,
         after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "ASCO 2026 Advances", size=32, bold=True, color=NAVY, after=5,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Clinical Impact, Comparative Treatment Value, Cost, Safety and Access",
         size=16.5, color=TEAL, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "A structured assessment of research and technology highlighted in recent ASCO and JCO email communications",
         size=12, italic=True, color=MID_GRAY, after=24,
         align=WD_ALIGN_PARAGRAPH.CENTER)

add_table(
    doc,
    ["Scope", "Assessment basis"],
    [
        ("Clinical evidence", "Primary ASCO/JCO reports and trial publications"),
        ("Comparators", "Established or current standard treatment in the same setting"),
        ("Impact measures", "PFS, RFS, OS, response, toxicity, burden, cost and access"),
        ("Cost perspective", "US/global references plus Indian MRP, retail or import references where available"),
        ("Cutoff", "June 9, 2026"),
    ],
    widths=[1.55, 4.8],
    font_size=9.7,
)
add_callout(
    doc,
    "Interpretation warning",
    "This report is a strategic evidence review, not a treatment guideline. Investigational therapies are explicitly labeled. Cross-trial comparisons are directional because populations, endpoints and follow-up differ. Costs exclude rebates, insurance design, administration, monitoring, hospitalization and management of adverse events.",
    fill="FFF1D6",
    title_color=GOLD,
)

doc.add_page_break()

add_heading(doc, "Executive Summary", 1)
add_para(
    doc,
    "The recent ASCO/JCO communications show oncology advancing through four linked mechanisms: treatment intensification where durable control is inadequate; treatment de-escalation where similar efficacy can be achieved with less toxicity; biomarker-defined therapy that concentrates benefit in smaller populations; and digital infrastructure that improves detection, trial access and operational reliability.",
)

add_callout(
    doc,
    "Overall conclusion",
    "The strongest near-term practice signals are tucatinib added to first-line HER2 maintenance, carboplatin omission in selected neoadjuvant HER2-positive breast cancer, first-line pemigatinib for FGFR2-rearranged cholangiocarcinoma, and durable personalized neoantigen therapy results in melanoma. The highest uncertainty remains around early-phase ADCs, neoadjuvant sacituzumab in bladder cancer and AI systems that have not yet demonstrated prospective improvements in survival, enrollment or total cost of care.",
)

impact_rows = [
    ("neoCARHP: omit carboplatin", "Roche + generic manufacturers", "Phase III", "Similar pCR; less hematologic toxicity", "Cost-saving", "High, pending survival"),
    ("HER2CLIMB-05: add tucatinib", "Pfizer/Seagen; Indian generics", "Phase III", "+8.6 months median PFS", "High added cost", "High"),
    ("GeparNuevo: durvalumab", "AstraZeneca", "Phase II", "OS HR 0.33 at 86.4 months", "High", "Moderate-high"),
    ("Intismeran + pembrolizumab", "Moderna + Merck/MSD", "Phase IIb", "5-year RFS 68.8% vs 49.1%", "Unknown", "Potentially transformative"),
    ("FIGHT-302 pemigatinib", "Incyte", "Phase III", "PFS 8.3 vs 6.8 months; ORR 47% vs 15%", "Very high", "High in FGFR2 subset"),
    ("OrigAMI-1 amivantamab", "Johnson & Johnson", "Phase Ib/II", "ORR 19%-29%; right-sided 22%", "Very high", "Promising, not practice-changing"),
    ("Sigvotatug vedotin", "Pfizer/Seagen", "Phase I", "ORR 29%; PFS 6.4 months in subgroup", "Unknown", "Promising"),
    ("SURE-01 sacituzumab", "Gilead Sciences", "Phase II", "ypT0N0-x 29.5%; 24-month EFS 71.4%", "High", "Exploratory"),
    ("AI trial matching", "Academic/institutional system", "Retrospective", "Sensitivity up to 98.7%; NPV 99.7%", "Implementation", "Operationally high; outcomes unproven"),
]
add_table(
    doc,
    ["Advance", "Company / manufacturer", "Maturity", "Measured value", "Cost", "Impact judgment"],
    impact_rows,
    widths=[1.15, 1.05, 0.65, 1.55, 0.75, 1.20],
    font_size=8.1,
)

add_heading(doc, "Impact Rating Categories", 1)
add_para(
    doc,
    "These ratings describe the strength and practical meaning of the evidence. They do not replace regulatory approval, professional guidelines or individual clinical judgment.",
)
rating_rows = [
    (
        "Exploratory",
        "An early signal that generates a hypothesis but is not reliable enough to guide routine treatment.",
        "Usually preclinical, retrospective, phase I, small single-arm phase II, biomarker subgroup or post hoc evidence.",
        "Do not infer superiority. Use mainly for research planning or trial referral.",
    ),
    (
        "Promising",
        "The treatment or technology shows a clinically interesting activity or operational performance signal that justifies further study.",
        "The observed effect appears meaningful, but evidence may be nonrandomized, based on a small subgroup, use a historical comparator or lack mature survival data.",
        "Reasonable for clinical trials; insufficient by itself for broad routine adoption.",
    ),
    (
        "Promising, not practice-changing yet",
        "The signal is stronger or more relevant than a purely exploratory result, but one or more decisive adoption requirements remain unmet.",
        "Common missing elements include randomized comparison with current care, phase III confirmation, overall-survival evidence, reproducible biomarker selection, regulatory approval, acceptable cost or real-world feasibility.",
        "Clinicians should not replace established standard care solely on this evidence.",
    ),
    (
        "Moderate impact",
        "Evidence supports a meaningful benefit, but the magnitude, population, safety, cost or evidence maturity limits broad applicability.",
        "May include positive randomized phase II evidence, a modest phase III benefit, or a strong result confined to a narrow population.",
        "Potential selective adoption after guidelines, approval and patient-level assessment.",
    ),
    (
        "High impact",
        "Robust evidence demonstrates a clinically meaningful advantage over a relevant current-care comparator.",
        "Usually a positive randomized phase III trial with meaningful absolute or relative benefit, manageable toxicity and a clear eligible population.",
        "Likely to influence guidelines or routine care, subject to approval, affordability and system readiness.",
    ),
    (
        "High impact, pending survival",
        "The treatment has strong evidence on an accepted intermediate endpoint but definitive long-term outcomes remain incomplete.",
        "Examples include noninferior pCR with substantially lower toxicity or a major PFS gain while OS remains immature.",
        "Adoption may begin selectively, but long-term follow-up can still change the benefit-risk judgment.",
    ),
    (
        "Potentially transformative",
        "The approach could materially change treatment strategy or create a new therapeutic platform if the result is confirmed.",
        "The observed benefit is unusually large, durable or mechanistically novel, but evidence is not yet sufficient for a definitive transformative claim.",
        "High strategic importance; requires confirmatory trials, scalable delivery, safety validation and an acceptable economic model.",
    ),
    (
        "Operationally high, clinical impact unproven",
        "The tool performs its operational task well, but no evidence yet shows that it improves patient outcomes.",
        "Typical for AI, workflow or diagnostic systems validated on accuracy, speed or agreement rather than enrollment, treatment, survival, equity or total cost.",
        "Use through monitored implementation with prospective outcome and subgroup measurement.",
    ),
]
add_table(
    doc,
    ["Category", "Meaning", "Evidence pattern", "Clinical implication"],
    rating_rows,
    widths=[1.25, 1.75, 1.8, 1.55],
    font_size=8.5,
)
add_callout(
    doc,
    "Meaning of “promising” in this report",
    "“Promising” means that the result is sufficiently encouraging to justify a larger or better-controlled study. It does not mean the treatment is proven to be better than current care, approved for that use, cost-effective or ready for routine clinical adoption.",
    fill="FFF1D6",
    title_color=GOLD,
)

add_heading(doc, "What Is Actually Advancing?", 1)
for item in [
    "Precision is moving from selecting a drug by one biomarker to integrating histology, genomics, transcriptomics, circulating DNA and computational pathology.",
    "Antibody-drug conjugates are expanding beyond metastatic rescue therapy into neoadjuvant and earlier-line treatment, but marrow toxicity, neuropathy, interstitial lung disease and price remain major constraints.",
    "Immunotherapy is being optimized by duration and sequencing, not only by adding more therapy. GeparNuevo raises the possibility that neoadjuvant-only checkpoint blockade may deliver durable benefit in some settings.",
    "De-escalation is becoming a measurable innovation. Removing carboplatin while maintaining pCR could improve completion rates, reduce supportive care and lower direct and indirect costs.",
    "AI is already used for information synthesis and data work, but clinical adoption is ahead of governance and training. Prospective outcome evidence remains limited.",
]:
    add_bullet(doc, item)

add_heading(doc, "Method and Interpretation Framework", 1)
add_para(doc, "Each treatment or technology was assessed using the following hierarchy:")
for text in [
    "Clinical benefit: absolute and relative change in survival, progression, recurrence or response.",
    "Evidence maturity: randomized phase III evidence carries more weight than single-arm or retrospective validation.",
    "Safety: grade 3-4 events, discontinuation, immune toxicity, marrow toxicity, organ-specific risks and long-term uncertainty.",
    "Treatment burden: oral versus infusion delivery, frequency, monitoring, surgery, manufacturing delay and supportive care.",
    "Economic exposure: drug acquisition, administration, biomarker testing and adverse-event management.",
    "Equity and scalability: dependence on high-complexity pathology, sequencing, manufacturing or digital infrastructure.",
]:
    add_number(doc, text)

add_heading(doc, "Detailed Treatment Assessments", 1)

add_profile(
    doc, 1,
    "Carboplatin De-escalation in HER2-Positive Early Breast Cancer (neoCARHP)",
    "Randomized noninferiority phase III; survival follow-up is not yet mature.",
    "Stage II-III HER2-positive breast cancer commonly receives a taxane, carboplatin, trastuzumab and pertuzumab before surgery. Carboplatin contributes cytotoxic activity but increases neutropenia, leukopenia, thrombocytopenia, anemia, nausea and treatment interruption.",
    "766 patients were included in the modified intention-to-treat analysis: taxane plus trastuzumab/pertuzumab without carboplatin (THP, n=382) versus with carboplatin (TCbHP, n=384). The primary endpoint was pathologic complete response and the trial was designed for noninferiority.",
    "pCR was 64.1% with THP and 65.9% with TCbHP; absolute difference -1.8 percentage points (95% CI -8.5 to 5.0), meeting the prespecified noninferiority test. Grade 3-4 neutropenia was 6.8% versus 16.4%, leukopenia 5.5% versus 14.8%, and diarrhea 2.6% versus 4.2%.",
    "The immediate value is not superior tumor eradication but similar pCR with materially lower acute toxicity. For every 100 treated patients, omitting carboplatin was associated with roughly 10 fewer grade 3-4 neutropenia events and 9 fewer severe leukopenia events, while pCR differed by less than 2 patients.",
    "Expected advantages include fewer transfusions, growth-factor injections, infections, dose delays and hospital visits. The tradeoff is uncertainty about event-free and overall survival, especially in biologically high-risk subgroups.",
    "Carboplatin is inexpensive as a generic drug, so pharmacy savings alone are modest. The larger economic gain is likely from avoided supportive care, laboratory monitoring, emergency visits and lost work. Trastuzumab/pertuzumab remain the major drug-cost drivers.",
    "High potential for a value-based change in selected patients because it improves tolerability and system capacity without obvious short-term efficacy loss.",
    "pCR is a surrogate, not definitive survival evidence. Generalizability outside the trial population and subgroup selection require longer follow-up. Until then, omission should not be interpreted as universally preferable.",
    "[1]",
    manufacturer="Roche supplies the originator HER2 antibodies trastuzumab and pertuzumab; taxanes and carboplatin are produced by multiple generic manufacturers.",
    india_access="All regimen components are commercially available. Illustrative Indian prices include carboplatin 450 mg at approximately ₹1,900-₹2,800 per vial, trastuzumab 440 mg biosimilars at roughly ₹17,700-₹57,500 per vial, and Perjeta 420 mg at approximately ₹2.52-₹2.62 lakh per vial. The carboplatin-omission strategy is a treatment approach, not a separate product launch. [19-22]",
)

add_profile(
    doc, 2,
    "Tucatinib Added to First-Line HER2 Maintenance (HER2CLIMB-05)",
    "Randomized double-blind phase III; PFS positive, overall survival immature.",
    "After induction taxane plus trastuzumab/pertuzumab for HER2-positive metastatic breast cancer, maintenance trastuzumab/pertuzumab is an established approach. Progression commonly occurs within approximately two years, and CNS relapse remains a major concern.",
    "654 women without progression after induction were randomized to tucatinib 300 mg twice daily or placebo, both with trastuzumab and pertuzumab. Twelve percent had current or previous brain metastases.",
    "Median investigator-assessed PFS was 24.9 months with tucatinib versus 16.3 months with placebo: an 8.6-month gain and 35.9% relative reduction in progression or death (HR 0.641; 95% CI 0.514-0.799). Blinded review was consistent. Benefit appeared across hormone-receptor and brain-metastasis subgroups. OS is immature.",
    "This is a direct randomized improvement over current maintenance care, unlike many cross-trial signals. The magnitude is clinically meaningful, but it adds continuous therapy rather than replacing an existing component.",
    "Diarrhea occurred in 72.7%; grade >=3 diarrhea in 6.1%. ALT and AST elevations occurred in 28.2% and 25.8%, with grade >=3 rates of 13.5% and 7.1%. Tucatinib was discontinued for treatment-emergent toxicity in 13.5%. This requires diarrhea management, liver monitoring and adherence support.",
    "The regimen adds an oral branded agent to two costly HER2 antibodies. Even before administration and monitoring, trastuzumab plus pertuzumab can exceed $100,000 annually at reference-product prices; tucatinib adds substantial pharmacy cost. Economic value will depend on OS, CNS-event avoidance and duration of treatment.",
    "High near-term practice impact because the comparator is current standard care and the PFS gain is large. The strongest value case may be in patients at elevated CNS risk if CNS-PFS and OS mature favorably.",
    "No mature OS or full cost-effectiveness result. Continuous toxicity and financial exposure must be weighed against the PFS gain. Regulatory status and guideline placement should be checked before routine adoption.",
    "[2, 15, 16]",
    manufacturer="Tucatinib was developed by Seagen and is now in Pfizer's oncology portfolio. In India, CDSCO-approved tucatinib is manufactured/marketed by MSN Laboratories, with additional Indian brands listed by other manufacturers.",
    india_access="CDSCO approved tucatinib 50 mg and 150 mg on April 8, 2025 for previously treated HER2-positive metastatic breast cancer with trastuzumab and capecitabine. A 150-mg 10-tablet strip is listed around ₹6,000 retail and ₹7,678-₹8,190 MRP; at 300 mg twice daily, drug-only exposure is roughly ₹72,000-₹98,000 per 30 days depending on brand and price. The HER2CLIMB-05 first-line maintenance use is not the same as the approved Indian indication and requires a separate regulatory/guideline assessment. [23-25]",
)

add_profile(
    doc, 3,
    "Neoadjuvant Durvalumab in Early Triple-Negative Breast Cancer (GeparNuevo)",
    "Randomized phase II with long-term follow-up; not the current approved TNBC standard.",
    "Current US standard for high-risk early TNBC is perioperative pembrolizumab with multiagent chemotherapy, including adjuvant continuation after surgery. This improves event-free and overall survival but creates prolonged immune exposure and high treatment cost.",
    "174 patients received durvalumab or placebo with neoadjuvant nab-paclitaxel followed by dose-dense epirubicin/cyclophosphamide. Durvalumab was not continued after surgery. Median follow-up was 86.4 months.",
    "Durvalumab improved invasive disease-free survival (HR 0.56), distant disease-free survival (HR 0.41) and overall survival (HR 0.33). The effect was seen despite the original pCR endpoint not showing a decisive overall improvement. Patients with baseline nodal disease appeared to derive greater iDFS benefit.",
    "The finding suggests durable immune priming may be achievable without routine adjuvant continuation, potentially reducing cumulative treatment. However, it is not a head-to-head comparison with pembrolizumab-based KEYNOTE-522 treatment and used a smaller phase II population.",
    "Checkpoint inhibitors can cause permanent endocrinopathies and serious immune-mediated pneumonitis, hepatitis, colitis and other organ toxicity. A neoadjuvant-only model could reduce exposure, but the study was not designed to prove that shorter therapy is equally safe or effective to current perioperative pembrolizumab.",
    "Pembrolizumab list price is approximately $12,272 per 200 mg dose every three weeks. Avoiding adjuvant doses could remove more than $100,000 in gross drug acquisition per patient, before administration, if equivalent outcomes were eventually proven. That equivalence has not been established.",
    "Scientifically important because it challenges treatment-duration assumptions and shows that pCR alone may not capture immune benefit.",
    "Small phase II trial, exploratory subgroup findings and no direct comparison with current standard pembrolizumab. Practice change requires a randomized duration/de-escalation trial.",
    "[3, 14]",
    manufacturer="AstraZeneca manufactures and markets durvalumab as Imfinzi.",
    india_access="Imfinzi is commercially available in India. The revised government-cited MRP for 500 mg/10 mL is ₹1,71,324 per vial, with pharmacy listings near ₹1.47-₹1.71 lakh. Neoadjuvant durvalumab for early TNBC as studied in GeparNuevo is not an established Indian labeled use; no separate launch date has been announced for this indication. [26-28]",
)

add_profile(
    doc, 4,
    "Personalized mRNA Neoantigen Therapy in Resected Melanoma (KEYNOTE-942)",
    "Randomized phase IIb; phase III confirmation ongoing; investigational.",
    "Adjuvant pembrolizumab reduces recurrence risk after resection of high-risk melanoma, but a substantial proportion of patients still relapse. Standard therapy is not individualized to each tumor's neoantigen set.",
    "157 patients with resected stage IIIB-IV melanoma were randomized 2:1 to nine intramuscular doses of individualized intismeran plus 18 pembrolizumab doses or pembrolizumab alone. The therapy is manufactured from patient-specific tumor sequencing.",
    "At median planned follow-up of 60.3 months, RFS HR was 0.51. Five-year RFS was 68.8% versus 49.1%, an absolute difference of 19.7 percentage points. DMFS HR was 0.41. Five-year OS was 92.2% versus 71.3%, but OS was exploratory and imprecise (HR 0.47; 95% CI 0.17-1.35).",
    "The absolute five-year recurrence-free difference is large, but the control group was small (n=50) and the study was not powered for definitive OS. Unlike adding another off-the-shelf drug, this approach creates patient-specific immune targets and may establish a new treatment platform.",
    "No new safety signal was reported; prior analyses found mainly injection-related systemic symptoms without clear amplification of pembrolizumab immune toxicity. Manufacturing time, tissue adequacy and sequencing quality are additional nonmedical burdens.",
    "No commercial price exists. Expected cost drivers include tumor sequencing, algorithmic antigen selection, individualized GMP manufacturing, logistics and pembrolizumab. Unit costs may initially exceed conventional biologics, although platform automation could reduce them at scale.",
    "Potentially transformative if phase III confirms the nearly 20-point absolute RFS difference. The broader impact is a reusable personalized-vaccine platform across tumor types.",
    "Phase IIb sample size, descriptive five-year analysis, immature/uncertain OS and no established manufacturing-cost model. Access may concentrate in advanced centers.",
    "[4]",
    manufacturer="Moderna and Merck/MSD jointly develop and plan to commercialize intismeran autogene; costs and profits are shared under their worldwide collaboration.",
    india_access="Not approved or commercially launched in India or elsewhere. There is no commercial price. The melanoma phase III INTerpath-001 study is fully enrolled, but neither company has announced an India launch date. A late-2020s launch is only a scenario and would depend on positive phase III data, global filing, individualized-manufacturing validation and Indian regulatory review. [29-31]",
)

add_profile(
    doc, 5,
    "First-Line Pemigatinib for FGFR2-Rearranged Cholangiocarcinoma (FIGHT-302)",
    "Randomized phase III; trial closed early after the standard-care landscape changed.",
    "Gemcitabine/cisplatin-based therapy, now commonly combined with immunotherapy in many jurisdictions, is the broad first-line standard. FGFR2 rearrangements define a molecular subgroup for which pemigatinib was previously used after progression.",
    "4,563 patients were prescreened, illustrating the large testing funnel. Only 167 were randomized to pemigatinib or gemcitabine/cisplatin before early closure. Crossover to pemigatinib was permitted.",
    "Median PFS was 8.3 versus 6.8 months (HR 0.58). ORR was 47% versus 15%, and median response duration 14.2 versus 6.3 months. Median OS was similar: 24.4 versus 25.0 months, influenced by crossover and limited sample size.",
    "Pemigatinib delivered a much higher and more durable response rate and a 1.5-month median PFS gain over chemotherapy, but no observed OS advantage. The comparator does not reflect every current chemo-immunotherapy regimen.",
    "FGFR inhibitors commonly cause hyperphosphatemia, nail/skin changes, stomatitis, diarrhea, fatigue and ocular toxicity, requiring phosphate and ophthalmologic monitoring. They avoid cisplatin nephrotoxicity, neuropathy, nausea and infusion burden.",
    "Manufacturer-reported 2025 WAC was $20,286 per 14-tablet bottle. At the standard 14-days-on/7-days-off schedule, gross annual acquisition is roughly $350,000 if treatment continues for a year. Biomarker testing and monitoring add cost, while oral delivery can reduce infusion burden.",
    "High impact within a small biomarker-defined population. It reinforces universal early molecular testing in cholangiocarcinoma and can provide rapid tumor shrinkage where chemotherapy response is limited.",
    "Small randomized sample, early closure, crossover, no OS improvement and an evolving comparator. Value depends strongly on sequencing relative to chemo-immunotherapy and negotiated price.",
    "[5, 13]",
    manufacturer="Incyte developed and markets pemigatinib as Pemazyre in approved markets.",
    india_access="No verified domestic commercial launch or Indian MRP was identified. Named-patient import may be possible, but this is not routine Indian market availability. The US list price of $20,286 per 14-tablet bottle converts to approximately ₹19.4 lakh at ₹95.4/US$, or about ₹3.34 crore for one year at the labeled 14-days-on/7-days-off schedule. This is a currency conversion, not an Indian selling price. No public India launch date is available, and first-line use remains indication-specific and evidence-dependent. [32-34]",
)

add_profile(
    doc, 6,
    "Amivantamab for Chemorefractory RAS/BRAF Wild-Type Colorectal Cancer (OrigAMI-1)",
    "Open-label phase Ib/II; investigational in colorectal cancer.",
    "Late-line metastatic colorectal cancer options include trifluridine/tipiracil-based therapy, regorafenib, fruquintinib and anti-EGFR rechallenge in selected molecularly eligible patients. Response rates after multiple lines are generally low; anti-EGFR rechallenge studies often report ORRs of approximately 8%-22%.",
    "94 patients with 2-3 prior metastatic lines and centrally confirmed RAS/BRAF/EGFR ectodomain wild-type, HER2-nonamplified disease received the EGFR-MET bispecific antibody. Cohorts were defined by tumor side and previous anti-EGFR exposure.",
    "ORR was 29% in left-sided anti-EGFR-naive disease, 19% after prior anti-EGFR treatment and 22% in right-sided disease. Median response duration was 6.1-9.8 months and median PFS 3.7-5.7 months.",
    "The signal is most notable in right-sided disease and after prior anti-EGFR therapy, settings where conventional EGFR antibodies are less reliable. However, there is no randomized comparison with current late-line treatments.",
    "The most frequent grade >=3 treatment-related events were rash 7%, acneiform dermatitis 4% and hypoalbuminemia 4%. Infusion reactions and skin toxicity are operational concerns, although only one patient discontinued for a treatment-related event.",
    "At a current cash reference of about $3,894 per 350-mg vial, a 1,050-mg dose uses three vials, approximately $11,700 per dose before administration. Every-two-week maintenance implies gross annual drug exposure near $300,000, excluding loading doses and wastage.",
    "Moderate potential if randomized trials confirm activity, particularly for right-sided disease and molecularly selected rechallenge. The dual EGFR-MET mechanism may overcome some resistance patterns.",
    "Small nonrandomized cohorts, biomarker-intensive selection, no survival comparison and very high acquisition cost. Phase III results are required.",
    "[6, 17]",
    manufacturer="Johnson & Johnson/Janssen develops and markets amivantamab as Rybrevant.",
    india_access="For colorectal cancer, amivantamab remains investigational and has no Indian launch date. Indian sites have participated in amivantamab clinical trials, and imported 350-mg vials are advertised by specialist suppliers, but no reliable official Indian MRP for the CRC setting was found. The US cash reference of about $3,894 per vial converts to about ₹3.72 lakh per vial at ₹95.4/US$; this is not an Indian market price. [35-37]",
)

add_profile(
    doc, 7,
    "Integrin Beta-6 ADC for Advanced NSCLC (Sigvotatug Vedotin)",
    "First-in-human phase I; randomized phase III comparison with docetaxel ongoing.",
    "After immunotherapy and platinum treatment, docetaxel-based therapy remains a common option for advanced NSCLC without a targetable alteration, but response and durability are limited and hematologic toxicity is substantial.",
    "117 patients were treated across dose-expansion groups. A clinically relevant subgroup included 42 taxane-naive patients with nonsquamous NSCLC. Sigvotatug targets integrin beta-6 and delivers the microtubule toxin MMAE.",
    "Across all patients, ORR was 19%, median response duration 11.3 months, PFS 3.6 months and OS 10.8 months. In taxane-naive nonsquamous disease, ORR was 29%, response duration 12.8 months, PFS 6.4 months and OS 14.8 months.",
    "The taxane-naive subgroup signal appears favorable to historical docetaxel expectations, but this is not a randomized comparison and patient selection may account for part of the difference.",
    "Peripheral sensory neuropathy was a key cumulative toxicity and caused several permanent discontinuations. At the selected schedule, 19% discontinued for treatment-emergent events and 10% discontinued because of sensory neuropathy.",
    "No commercial price exists. If approved, an ADC is likely to have high acquisition cost plus infusion and neuropathy-monitoring costs. Economic value would require better survival or lower severe toxicity than generic docetaxel.",
    "Potentially meaningful for biomarker-poor NSCLC if phase III confirms durable response and survival. The long response duration among responders is the strongest signal.",
    "Phase I design, heterogeneous dosing and populations, historical comparator and no validated selection biomarker in routine care.",
    "[7]",
    manufacturer="Originally developed by Seagen and now part of Pfizer's oncology pipeline following Pfizer's acquisition of Seagen.",
    india_access="Investigational globally; not approved or launched in India and no commercial price exists. A randomized phase III comparison with docetaxel is ongoing. Pfizer has described the product as a potential future launch but has not published an India date. A late-2020s India launch would be contingent on positive phase III results and subsequent global and CDSCO approvals. [38-40]",
)

add_profile(
    doc, 8,
    "Neoadjuvant Sacituzumab Govitecan in Cisplatin-Ineligible MIBC (SURE-01)",
    "Single-arm phase II; investigational and now compared against a rapidly changing standard.",
    "For cisplatin-ineligible patients, cystectomy alone historically produced substantial recurrence. Perioperative enfortumab vedotin plus pembrolizumab has now demonstrated approximately 57% pCR and survival improvement and is becoming the stronger benchmark.",
    "44 patients with cT2-T4aN0M0 disease received four cycles of sacituzumab before planned cystectomy. The starting dose was reduced from 10 to 7.5 mg/kg with primary neutropenia prophylaxis after two early deaths, one treatment-related.",
    "Protocol-defined ypT0N0 after cystectomy was 9.1%. Including repeat transurethral resection in patients who declined cystectomy, ypT0N0-x was 29.5%. Twenty-four-month EFS was 71.4%. Nonluminal tumors had more ypT0 responses than luminal tumors (46% versus 14%).",
    "Activity was demonstrated, but the pCR benchmark is lower than the approximately 57% reported with perioperative enfortumab/pembrolizumab. Surgery refusal and alternative response assessment complicate interpretation.",
    "After dose amendment, grade 3-4 treatment-related events occurred in 13.9%. Neutropenia, diarrhea and infection risk are central sacituzumab concerns. One treatment-related death in a 44-patient study is clinically important.",
    "The standard 10-mg/kg metastatic regimen is a high-cost infusion on days 1 and 8 every 21 days; the neoadjuvant reduced dose may lower drug use but adds G-CSF, infusion and laboratory costs. No value case is established against enfortumab/pembrolizumab.",
    "Useful as proof that TROP2 is active in MIBC and that molecular subtype/TOP1 may guide future ADC selection. It is unlikely to displace the current leading perioperative regimen without comparative evidence.",
    "Small single-center-style phase II evidence, protocol changes, missing cystectomy assessments, short follow-up and a superior contemporary benchmark.",
    "[8]",
    manufacturer="Gilead Sciences markets sacituzumab govitecan as Trodelvy after acquiring Immunomedics.",
    india_access="Trodelvy is not established in a domestic Indian commercial supply chain for this use; named-patient import is advertised. Import listings are approximately ₹1.85-₹1.90 lakh per 180-200 mg vial, but these are supplier quotes rather than an official Indian MRP. The neoadjuvant bladder-cancer use in SURE-01 is investigational, and no India launch date for this indication has been announced. [41-43]",
)

doc.add_page_break()
add_heading(doc, "Technology and Care-Delivery Advances", 1)

add_heading(doc, "AI-Enabled Clinical-Trial Matching", 2)
add_para(doc, "TRIAGE used full longitudinal electronic health records and versioned trial protocols to automate eligibility screening. In a retrospectively adjudicated sample, the predefined threshold produced 78.3% sensitivity, 98.5% specificity, 92.6% positive predictive value and 94.8% negative predictive value. A lower threshold increased sensitivity to 98.7% with 97.6% specificity and 99.7% NPV. Criterion-level agreement was 94.2% after structured readjudication.")
add_para(doc, "Compared with manual coordinator screening, the likely impact is faster review of very large records, more consistent criterion application and fewer missed trials. The tool can cite source evidence, which is important for auditability. However, retrospective accuracy does not prove that enrollment increases, screening time falls, disparities narrow or survival improves.")

add_heading(doc, "AI Adoption and Workforce Readiness", 2)
add_para(doc, "A survey of 227 oncology faculty and trainees found that 58% used AI several times per month or more, while 15% had never used it. Seventy-four percent expected AI to improve cancer diagnosis within ten years, but 93% wanted dedicated training and roughly half did not know where to find reliable learning resources. Female respondents and people age 60 or older reported lower frequent use after adjustment.")
add_para(doc, "The implementation risk is therefore not only algorithm accuracy. Unequal adoption, weak governance, privacy failures and automation bias can create new care gaps. Institutions should measure error rates by patient subgroup, require traceable source evidence and preserve clinician accountability.")

add_heading(doc, "Machine Learning for Earlier Monoclonal Gammopathy Detection", 2)
add_para(doc, "Using routine laboratory trajectories, a seven-variable XGBoost model predicted monoclonal protein within five years with AUC 0.84, positive predictive value 0.62 and negative predictive value 0.89. The model used age, lymphocyte trajectory, RBC count, total protein, RDW, BUN and eosinophils.")
add_para(doc, "The potential benefit is low-cost case finding from tests already collected, allowing targeted electrophoresis rather than universal screening. The major risk is false-positive testing and downstream biopsies. Prospective calibration in diverse populations and demonstration of improved clinical outcomes are required.")

add_heading(doc, "Computational Pathology, ctDNA and Molecular Infrastructure", 2)
add_para(doc, "The email portfolio also highlights computational HER2 quantification, standardized ctDNA assay validation, cell-free DNA toxicity monitoring and unified genotyping strategies. These technologies can reduce subjective pathology interpretation, reveal emerging resistance and improve trial selection. Their impact depends on analytical validity, turnaround time, specimen quality, interoperability and whether the result changes treatment.")

add_table(
    doc,
    ["Technology", "Immediate operational gain", "Clinical endpoint still needed", "Main cost/risk"],
    [
        ("AI trial matching", "Faster prescreening; fewer missed protocols", "Enrollment, time-to-enrollment, survival, equity", "Integration, governance, false eligibility"),
        ("Routine-lab ML", "Low-cost targeted case finding", "Prospective detection benefit and downstream harm", "False positives and extra testing"),
        ("Computational pathology", "More quantitative biomarker scoring", "Better treatment selection and outcomes", "Digital pathology infrastructure"),
        ("ctDNA/MRD", "Earlier molecular response/resistance signal", "Actionable intervention improves survival", "Assay cost, false reassurance, unclear action"),
        ("Personalized neoantigen manufacturing", "Patient-specific immune targeting", "Phase III RFS/OS and manufacturing reliability", "High complexity and access concentration"),
    ],
    widths=[1.25, 1.65, 2.15, 1.3],
    font_size=8.9,
)

add_heading(doc, "Global Growth, Equity and System Capacity", 1)
add_para(doc, "ASCO's global-oncology publications show that therapeutic innovation cannot be separated from diagnostic timing, treatment availability and registry quality.")

add_heading(doc, "Brazil: Prostate Cancer Treatment Inequity", 2)
add_para(doc, "A national public-system analysis included 670,205 men treated from 2008 to March 2023. Metastatic disease at diagnosis was recorded in 21.2% of Black/Pardo patients versus 19.7% of White patients. White patients received more local therapy at stage I, while non-White patients received more local therapy at stage III. Average recorded expenditure was 16.2% higher for White patients. Among 125,759 stage IV patients receiving systemic treatment, only 17.8% received docetaxel and approximately one quarter received older antiandrogens or estrogens without established survival benefit.")
add_para(doc, "Impact: universal coverage may narrow postdiagnosis access, but late detection, referral pathways and outdated systemic therapy can still produce unequal outcomes. New expensive treatments will widen rather than narrow gaps unless screening, staging and procurement improve.")

add_heading(doc, "Kyrgyzstan: Rising Recorded Cancer Burden", 2)
add_para(doc, "Reported new cases increased from 5,471 in 2020 to 6,651 in 2024, while registered prevalence reached 34,554. The study projected 7,306 new cases and prevalence of 35,772 in 2026. Stomach, lung, breast, cervical, liver and colorectal cancers made up the leading burden.")
add_para(doc, "Impact: rising incidence increases demand for pathology, surgery, radiotherapy, systemic therapy and survivorship services. However, 2025-2026 values were modeled projections and age-standardized estimates relied partly on scaled external baselines, so capacity planning should use ranges rather than a single forecast.")

doc.add_page_break()
add_heading(doc, "Cost and Value Assessment", 1)
add_callout(
    doc,
    "How to read the cost figures",
    "The figures below are gross list-price, MRP, retail-listing or import-reference figures, not net prices and not patient out-of-pocket amounts. They exclude rebates, patient-assistance programs, biosimilar discounts, administration, imaging, laboratory testing, supportive care, toxicity management and hospitalization. USD conversions use an illustrative June 2026 rate of ₹95.4 per US dollar. A converted US price is not an Indian market price.",
    fill="FFF1D6",
    title_color=GOLD,
)

cost_rows = [
    ("Pembrolizumab", "Merck/MSD", "$12,272 per 200-mg dose", "Keytruda 100 mg MRP ₹2.165 lakh; retail listing ₹1.948 lakh", "Commercially available; indication-specific approval applies"),
    ("Pemigatinib", "Incyte", "$20,286 per 14 tablets", "≈₹19.4 lakh converted; no verified Indian MRP", "No verified domestic launch; import/access inquiry only"),
    ("Amivantamab", "Johnson & Johnson", "$3,894 per 350-mg vial", "≈₹3.72 lakh converted; supplier quotes are not official MRP", "CRC use investigational; no India launch date"),
    ("Pertuzumab", "Roche", "$5,534 per 420-mg vial", "Perjeta 420 mg MRP about ₹2.52-₹2.62 lakh", "Commercially available"),
    ("Trastuzumab 440 mg", "Roche originator; multiple biosimilars", "US reference varies by product", "Indian biosimilar listings about ₹17,700-₹57,500; originator MRP ₹1.107 lakh", "Commercially available; wide brand-price variation"),
    ("Tucatinib 150 mg", "Pfizer/Seagen; Indian generics", "US branded price substantial", "About ₹6,000 retail/₹7,678-₹8,190 MRP per 10 tablets", "India approved April 8, 2025; HER2CLIMB-05 use differs"),
    ("Durvalumab 500 mg", "AstraZeneca", "US price varies by contract", "Imfinzi revised MRP ₹1.713 lakh; listings ₹1.47-₹1.71 lakh", "Commercially available; GeparNuevo use not established label"),
    ("Intismeran autogene", "Moderna + Merck/MSD", "No commercial price", "No price", "Investigational; no public India launch date"),
    ("Sigvotatug vedotin", "Pfizer/Seagen", "No commercial price", "No price", "Investigational; phase III ongoing"),
    ("Sacituzumab govitecan", "Gilead Sciences", "High-cost ADC; US price varies", "Import listings about ₹1.85-₹1.90 lakh per vial", "Imported access; SURE-01 use investigational"),
]
add_table(
    doc,
    ["Therapy", "Company", "US/global gross reference", "India price reference", "India availability / launch"],
    cost_rows,
    widths=[1.05, 1.05, 1.25, 1.55, 1.45],
    font_size=8.3,
)

add_heading(doc, "India Commercialization and Launch Outlook", 2)
india_rows = [
    ("neoCARHP components", "Multiple; Roche for HER2 antibodies", "Available", "No new product launch needed", "The de-escalation strategy still requires guideline and clinician adoption."),
    ("Tucatinib", "Pfizer/Seagen globally; MSN and other Indian manufacturers", "Approved and marketed", "Launched after April 8, 2025 approval", "Current Indian approval is not the same first-line maintenance setting as HER2CLIMB-05."),
    ("Durvalumab", "AstraZeneca", "Available", "Already launched", "GeparNuevo early-TNBC use remains off-label/investigational from an India-label perspective."),
    ("Intismeran autogene", "Moderna + Merck/MSD", "Not launched", "No announced date; late-2020s at earliest is an inference", "Requires positive phase III results, global approval and individualized manufacturing capacity."),
    ("Pemigatinib", "Incyte", "No verified domestic launch", "No public India date", "Imported access may be possible; first-line positioning needs indication-specific approval."),
    ("Amivantamab in CRC", "Johnson & Johnson", "Clinical-trial/import context", "No public India date for CRC", "Rybrevant approvals in other tumor settings do not establish CRC use."),
    ("Sigvotatug vedotin", "Pfizer/Seagen", "Not launched", "No announced date; late-2020s at earliest is an inference", "Phase III efficacy, safety and regulatory review are still required."),
    ("Sacituzumab in neoadjuvant MIBC", "Gilead Sciences", "Imported product; study use not launched", "No public India date for this indication", "The SURE-01 regimen requires comparative confirmation and regulatory review."),
]
add_table(
    doc,
    ["Advance", "Company / manufacturer", "India status", "Anticipated launch", "Qualification"],
    india_rows,
    widths=[1.1, 1.25, 1.05, 1.25, 1.70],
    font_size=8.3,
)

add_heading(doc, "Value by Type of Innovation", 2)
for item in [
    "De-escalation value: neoCARHP may reduce toxicity and total care cost even though carboplatin itself is inexpensive.",
    "Add-on value: HER2CLIMB-05 gains 8.6 months of median PFS but adds continuous drug cost and toxicity. OS and CNS-event reduction determine long-term cost-effectiveness.",
    "Biomarker value: pemigatinib concentrates benefit in FGFR2-rearranged disease, but requires broad testing of many patients to identify a small eligible group.",
    "Platform value: personalized neoantigen therapy may generate major recurrence reduction, but manufacturing economics and turnaround time are part of clinical value.",
    "Operational value: AI can lower screening labor and missed opportunities, but savings should be measured against software, integration, quality assurance and clinician-review costs.",
]:
    add_bullet(doc, item)

add_heading(doc, "Impact Measurement Scorecard", 1)
score_rows = [
    ("Clinical efficacy", "Absolute survival/PFS/RFS gain; NNT where calculable", "At least clinically meaningful and statistically credible"),
    ("Safety", "Grade 3-5 toxicity, discontinuation, irreversible toxicity", "Net reduction or justified increase"),
    ("Patient burden", "Infusion time, oral adherence, monitoring, travel, surgery", "Lower burden or clear benefit tradeoff"),
    ("Economic value", "Incremental drug + administration + AE cost per outcome", "Competitive ICER/budget impact"),
    ("Equity", "Testing availability, geography, race, income, age", "No widening of access gap"),
    ("Evidence maturity", "Phase, comparator, follow-up, external validity", "Randomized evidence for adoption"),
    ("System readiness", "Pathology, sequencing, pharmacy, digital integration", "Reliable turnaround and auditability"),
]
add_table(
    doc,
    ["Domain", "Measure", "Adoption expectation"],
    score_rows,
    widths=[1.35, 3.15, 1.85],
    font_size=9.3,
)

add_heading(doc, "Recommended Strategic Priorities", 1)
for text in [
    "Prioritize de-escalation pathways where noninferior efficacy and lower toxicity can be confirmed with survival follow-up.",
    "Require molecular testing early enough to influence first-line treatment, especially in cholangiocarcinoma and precision-oncology pathways.",
    "Build total-cost models that include toxicity, administration and patient time rather than comparing drug prices alone.",
    "Treat phase I/II ADC results as portfolio signals, not as established treatment superiority.",
    "Pilot AI with prospective workflow endpoints: coordinator time, enrollment, false exclusion, subgroup performance and cost per enrolled patient.",
    "Pair adoption of expensive precision therapies with equity measures: testing completion, turnaround time, denial rates and geographic access.",
]:
    add_number(doc, text)

add_heading(doc, "Final Assessment", 1)
add_para(doc, "The ASCO evidence set does not point to one universal breakthrough. It shows a portfolio transition toward more individualized, earlier and technologically enabled cancer care. The largest clinical gains may come from targeted or immune treatments, but the largest system gains may come from better selection, shorter treatment, fewer toxic events and faster access to trials.")
add_para(doc, "A treatment should be considered an advancement only when its benefit exceeds the combined burden of toxicity, cost, complexity and unequal access. On that basis, neoCARHP is a strong value signal; HER2CLIMB-05 and FIGHT-302 are strong efficacy signals with major cost implications; intismeran is a high-upside platform requiring phase III confirmation; and the ADC and AI studies are promising technologies whose real-world impact remains to be demonstrated.")

doc.add_page_break()
add_heading(doc, "Sources", 1)
sources = [
    "[1] Gao et al. neoCARHP randomized phase III. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-25-02176",
    "[2] Dieras et al. HER2CLIMB-05. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-25-02600",
    "[3] Loibl et al. GeparNuevo long-term analysis. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-25-02311",
    "[4] Khattak et al. KEYNOTE-942 five-year update. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-26-00835",
    "[5] Bekaii-Saab et al. FIGHT-302 pemigatinib. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-26-00788",
    "[6] Oberstein et al. OrigAMI-1 amivantamab in mCRC. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-25-02187",
    "[7] Peters et al. Sigvotatug vedotin phase I. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-25-02016",
    "[8] Necchi et al. SURE-01 sacituzumab in MIBC. J Clin Oncol. 2026. https://ascopubs.org/doi/10.1200/JCO-26-00142",
    "[9] Patel et al. AI-enabled trial matching. JCO Oncol Pract. 2026. https://ascopubs.org/doi/10.1200/OP-26-00076",
    "[10] Schadler et al. AI knowledge and use in oncology. JCO Clin Cancer Inform. 2026. https://ascopubs.org/doi/10.1200/CCI-25-00300",
    "[11] Movassagh et al. ML prediction of monoclonal gammopathy. JCO Clin Cancer Inform. 2026. https://ascopubs.org/doi/10.1200/CCI-26-00013",
    "[12] Herchenhorn et al. Prostate cancer inequities in Brazil. JCO Glob Oncol. 2026. https://ascopubs.org/doi/10.1200/GO-25-00712",
    "[13] Rehman et al. Cancer prevalence trends in Kyrgyzstan. JCO Glob Oncol. 2026. https://ascopubs.org/doi/10.1200/GO-26-00137",
    "[14] Merck Colorado WAC disclosure, Keytruda. 2026. https://www.merckconnect.com/wp-content/uploads/sites/110/2024/04/Colorado-Pricing-Sheets-KEYTRUDA-pembrolizumab.pdf",
    "[15] Incyte Colorado WAC disclosure, Pemazyre. 2025. https://cdn.incyte.com/Assets/Pemazyre-Colorado_pricing-data_04.04.2025.pdf",
    "[16] Perjeta cash-price reference. https://www.drugs.com/medical-answers/perjeta-cost-3550933/",
    "[17] Herceptin cash-price reference. https://www.drugs.com/price-guide/herceptin",
    "[18] Rybrevant cash-price reference. https://www.drugs.com/price-guide/rybrevant",
    "[19] PharmEasy. Carboplatin 450 mg Indian MRP/retail listing. https://pharmeasy.in/online-medicine-order/carboplatin-450mg-inj-45ml-39070",
    "[20] Tata 1mg. Indian trastuzumab 440 mg prices and biosimilar alternatives. https://www.1mg.com/drugs/priunta-440-injection-709856",
    "[21] Tata 1mg. Hertraz 440 mg Indian MRP and alternatives. https://www.1mg.com/drugs/Hertraz-Injection-336821",
    "[22] Netmeds and Macarius Health. Perjeta 420 mg Indian MRP references. https://www.netmeds.com/prescriptions/perjeta-420mg-injection-14ml",
    "[23] CDSCO. List of new drugs approved in 2025: tucatinib approval dated April 8, 2025. https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadApprovalNewDrugs/List%20of%20New%20Drugs%20approved%20258april.pdf",
    "[24] CDSCO Oncology SEC, April 22, 2025. Tucatinib manufacture and marketing recommendation for MSN Laboratories. https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadCommitteeFiles/Recommendations%20Oncology%2022.04.2025.pdf",
    "[25] Macarius Health. Tucaone 150 mg Indian price listing. https://www.macariushealth.com/products/oncology-cancer-care/tucaone-150mg-tablet-150mg/",
    "[26] Lok Sabha annexure. Revised Indian MRP for Imfinzi 120 mg and 500 mg. https://sansad.in/getFile/loksabhaquestions/annex/183/AU1952_mtQhcK.pdf?source=pqals",
    "[27] Tata 1mg. Imfinzi 500 mg Indian MRP and retail listing. https://www.1mg.com/drugs/imfinzi-injection-684073",
    "[28] CDSCO Oncology SEC, April 22, 2025. Durvalumab manufacturer and Indian indication review. https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadCommitteeFiles/Recommendations%20Oncology%2022.04.2025.pdf",
    "[29] Merck and Moderna. Intismeran five-year data and phase III program status, January 20, 2026. https://www.merck.com/news/moderna-merck-announce-5-year-data-for-intismeran-autogene-in-combination-with-keytruda-pembrolizumab-demonstrated-sustained-improvement-in-the-primary-endpoint-of-recurrence-free-survival-i/",
    "[30] Merck and Moderna. Joint development and commercialization agreement for V940. https://www.merck.com/news/merck-and-moderna-announce-exercise-of-option-by-merck-for-joint-development-and-commercialization-of-investigational-personalized-cancer-vaccine/",
    "[31] Merck pipeline. Intismeran autogene phase II/III development program. https://www.merck.com/research/product-pipeline/",
    "[32] Incyte. Pemazyre US WAC disclosure, April 2025. https://cdn.incyte.com/Assets/Pemazyre-Colorado_pricing-data_04.04.2025.pdf",
    "[33] Incyte/Pemazyre product information. https://www.pemazyre.com/",
    "[34] June 2026 USD/INR reference used for illustrative conversion. https://officialforexrates.com/",
    "[35] Johnson & Johnson. Rybrevant/amivantamab product and development information. https://www.jnj.com/innovativemedicine/oncology",
    "[36] CDSCO. India clinical-trial permission involving amivantamab. https://cdsco.gov.in/opencms/resources/UploadCDSCOWeb/2018/UploadCTApprovals/%2820%29-45349-JJ-CTNOC.pdf",
    "[37] GNH India. Imported Rybrevant access listing; no official Indian MRP stated. https://www.gnhindia.com/products/rybrevant-350mg-ge-vials",
    "[38] Pfizer. Completion of Seagen acquisition and oncology portfolio ownership. https://www.pfizer.com/news/press-release/press-release-detail/pfizer-completes-acquisition-seagen",
    "[39] Pfizer. Sigvotatug vedotin investigational development and potential-launch planning. https://www.pfizer.com/about/careers/job/4956005",
    "[40] ClinicalTrials.gov. Phase III sigvotatug vedotin versus docetaxel, NCT06012435. https://clinicaltrials.gov/study/NCT06012435",
    "[41] Gilead. Trodelvy prescribing information and manufacturer information. https://www.trodelvyhcp.com/-/media/project/trodelvy/trodelvyhcp/hcp/resource-pdf/trodelvy_pi.pdf",
    "[42] Indian Pharma Network. Named-patient import information for Trodelvy in India. https://www.indianpharmanetwork.in/trodelvy-sacituzumab-govitecan-hziy/",
    "[43] TradeIndia supplier listing. Illustrative imported Trodelvy price; not an official MRP. https://www.tradeindia.com/products/trodelvy-180mg-injection-10010514.html",
]
for source in sources:
    add_para(doc, source, size=9.3, color=DARK, after=5, line=1.06)

add_para(
    doc,
    "Prepared June 9, 2026. Clinical and pricing information changes over time; regulatory status, guidelines and current net prices should be rechecked before decision-making.",
    size=9.1, italic=True, color=MID_GRAY, before=10, after=0,
)

doc.save(OUT)
print(OUT)
